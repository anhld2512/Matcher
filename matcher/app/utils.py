import os
import shutil
from typing import Tuple

import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from docx import Document


def detect_file_type(file_path: str) -> str:
    """Return 'docx' or 'pdf' based on file extension (case‑insensitive)."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == ".docx":
        return "docx"
    if ext == ".pdf":
        return "pdf"
    raise ValueError(f"Unsupported file type: {ext}")


def extract_text_from_docx(file_path: str) -> str:
    """Read a .docx file and return its full text content including tables."""
    doc = Document(file_path)

    # Extract text from paragraphs
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def _extract_text_from_pdf_using_mupdf(file_path: str) -> str:
    """Attempt to extract text from a PDF using PyMuPDF. Returns empty string if none found."""
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text.strip()


def _extract_text_from_pdf_using_ocr(file_path: str) -> str:
    """Fallback OCR extraction: convert PDF pages to images and run pytesseract."""
    images = convert_from_path(file_path)
    ocr_text = ""
    for img in images:
        ocr_text += pytesseract.image_to_string(img) + "\n"
    return ocr_text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF.
    First tries PyMuPDF; if the result is empty, falls back to OCR.
    """
    text = _extract_text_from_pdf_using_mupdf(file_path)
    if text:
        return text
    # fallback to OCR
    return _extract_text_from_pdf_using_ocr(file_path)


def extract_text(file_path: str) -> str:
    """Unified entry point – detects file type and extracts text accordingly."""
    ftype = detect_file_type(file_path)
    if ftype == "docx":
        return extract_text_from_docx(file_path)
    if ftype == "pdf":
        return extract_text_from_pdf(file_path)
    raise ValueError("Unsupported file type")


def save_uploaded_file(uploaded_file, destination_dir: str) -> str:
    """Save an uploaded file (FastAPI UploadFile) to the destination directory.
    Returns the absolute path of the saved file.
    """
    os.makedirs(destination_dir, exist_ok=True)
    dest_path = os.path.join(destination_dir, uploaded_file.filename)
    with open(dest_path, "wb") as out_file:
        shutil.copyfileobj(uploaded_file.file, out_file)
    return dest_path
